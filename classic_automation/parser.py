# ============================================================
#  parser.py — извлекает параметры варианта из .doc/.docx
# ============================================================
import re
import subprocess
import tempfile
import os
from pathlib import Path
from dataclasses import dataclass, field
from typing import Optional


@dataclass
class VariantParams:
    """Параметры конкретного варианта задания."""
    variant:  int   = 0

    # W2(s) = K1 (усиление)
    K1: float = 1.0

    # W3(s) = K3 / (T3*s + 1)
    K3: float = 1.0
    T3: float = 1.0

    # W4(s) = K4 / (T4*s + 1)
    K4: float = 1.0
    T4: float = 1.0

    # W5(s) = K5 / s  (интегратор)
    K5: float = 1.0

    # Вычисляемые поля (заполняет calculator.py)
    WP_num: list = field(default_factory=list)
    WP_den: list = field(default_factory=list)


def _doc_to_text(doc_path: str) -> str:
    """Конвертирует .doc/.docx в plain-text через LibreOffice (или python-docx)."""
    path = Path(doc_path)
    suffix = path.suffix.lower()

    # Пробуем python-docx (только для .docx)
    if suffix == ".docx":
        try:
            from docx import Document
            doc = Document(str(path))
            return "\n".join(p.text for p in doc.paragraphs)
        except ImportError:
            pass

    # Fallback: LibreOffice --convert-to txt
    try:
        with tempfile.TemporaryDirectory() as tmp:
            cmd = ["soffice", "--headless", "--convert-to", "txt:Text", str(path), "--outdir", tmp]
            subprocess.run(cmd, capture_output=True, timeout=30)
            txt_file = Path(tmp) / (path.stem + ".txt")
            if txt_file.exists():
                return txt_file.read_text(encoding="utf-8", errors="ignore")
    except Exception:
        pass

    # Последний fallback: только .docx — бинарный .doc как текст даёт мусор
    if suffix != ".doc":
        return path.read_text(encoding="utf-8", errors="ignore")
    return ""


# ── таблица вариантов (пример для нескольких вариантов) ──────────────────────
# Формат: вариант → (K1, K3, T3, K4, T4, K5)
# Заполните полностью по своей таблице вариантов
VARIANT_TABLE = {
     1: (50,  2.0, 0.4, 0.15, 0.04, 0.04),
     2: (60,  2.2, 0.45, 0.18, 0.045, 0.045),
     3: (70,  2.3, 0.42, 0.16, 0.042, 0.042),
     4: (80,  2.4, 0.48, 0.17, 0.048, 0.048),
     5: (90,  2.1, 0.41, 0.14, 0.041, 0.041),
     6: (75,  2.5, 0.46, 0.19, 0.046, 0.046),
     7: (65,  2.0, 0.44, 0.16, 0.044, 0.044),
     8: (85,  2.3, 0.43, 0.15, 0.043, 0.043),
     9: (55,  2.2, 0.47, 0.18, 0.047, 0.047),
    10: (95,  2.4, 0.40, 0.14, 0.040, 0.040),
    11: (72,  2.1, 0.46, 0.17, 0.046, 0.046),
    12: (68,  2.5, 0.42, 0.16, 0.042, 0.042),
    13: (82,  2.0, 0.44, 0.15, 0.044, 0.044),
    14: (78,  2.3, 0.48, 0.19, 0.048, 0.048),
    15: (62,  2.2, 0.41, 0.14, 0.041, 0.041),
    16: (92,  2.4, 0.45, 0.17, 0.045, 0.045),
    17: (100, 2.5, 0.50, 0.20, 0.050, 0.050),  # ваш вариант
    18: (58,  2.1, 0.43, 0.16, 0.043, 0.043),
    19: (88,  2.3, 0.47, 0.18, 0.047, 0.047),
    20: (73,  2.5, 0.46, 0.15, 0.046, 0.046),
}


def parse_variant_number(doc_path: str) -> Optional[int]:
    """Ищет номер варианта в тексте документа."""
    text = _doc_to_text(doc_path)

    # Паттерны: "вариант 17", "v17", "LP_v17", "#17" и т.д.
    patterns = [
        r"[Вв]ариант\s*[№#]?\s*(\d{1,2})",
        r"[Vv]ariant\s*[№#]?\s*(\d{1,2})",
        r"LP_v(\d{1,2})",
        r"_v(\d{1,2})[_\.]",
        r"[№#]\s*(\d{1,2})",
    ]
    for pat in patterns:
        m = re.search(pat, text)
        if m:
            return int(m.group(1))

    # Попробуем извлечь из имени файла
    fname = Path(doc_path).stem
    m = re.search(r"[vV](\d{1,2})", fname)
    if m:
        return int(m.group(1))

    return None


def parse_variant(doc_path: str) -> VariantParams:
    """
    Главная функция парсинга.
    Возвращает VariantParams с заполненными параметрами блоков.
    """
    variant = parse_variant_number(doc_path)
    if variant is None:
        raise ValueError(f"Не удалось определить номер варианта в файле: {doc_path}")

    if variant not in VARIANT_TABLE:
        raise ValueError(
            f"Вариант {variant} не найден в VARIANT_TABLE. "
            f"Добавьте его в parser.py → VARIANT_TABLE."
        )

    K1, K3, T3, K4, T4, K5 = VARIANT_TABLE[variant]
    return VariantParams(
        variant=variant,
        K1=K1,
        K3=K3, T3=T3,
        K4=K4, T4=T4,
        K5=K5,
    )


def parse_q9_text(doc_path: str) -> str:
    """Возвращает текст вопроса 9 из .docx файла задания (без номера '9.')."""
    path = Path(doc_path)
    if path.suffix.lower() != '.docx':
        return ""
    try:
        from docx import Document
        doc = Document(str(path))
        for p in doc.paragraphs:
            t = p.text.strip()
            if re.match(r'^9[\.\s]', t) and 'задач' in t.lower():
                return re.sub(r'^9[\.\s]+', '', t).strip()
    except Exception:
        pass
    return ""


if __name__ == "__main__":
    import sys
    path = sys.argv[1] if len(sys.argv) > 1 else "LP_v17.doc"
    p = parse_variant(path)
    print(f"Вариант: {p.variant}")
    print(f"W2: K1={p.K1}")
    print(f"W3: {p.K3}/({p.T3}s+1)")
    print(f"W4: {p.K4}/({p.T4}s+1)")
    print(f"W5: {p.K5}/s")
