# report.py — заполняет задание LP_vXX.doc результатами и скриншотами
import re
import shutil
from pathlib import Path

from docx import Document
from docx.shared import Inches

from calculator import CalcResults
from classic_gui import Screenshots

ELLIPSIS = '…'  # U+2026
_COL_W = 10          # ширина колонки таблицы ПФ (9 пробелов + символ)
_BLK_ID_W  = 12      # ширина поля block-id в текстовой таблице блоков
_BLK_VAL_W = 10      # ширина числовых полей в таблице блоков
_BLK_CONN_W = 10     # ширина поля связей

_BLK_ROW_RE = re.compile(r'^\|    #(\d)      \|')


# ── Helpers для коэффициентных таблиц ПФ ─────────────────────────────────────

def _convert_doc_to_docx(doc_path: str, out_path: str) -> str:
    """Конвертирует .doc → .docx через Word COM."""
    import win32com.client
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    try:
        doc = word.Documents.Open(str(Path(doc_path).resolve()))
        doc.SaveAs2(str(Path(out_path).resolve()), FileFormat=16)
        doc.Close(False)
    finally:
        word.Quit()
    return out_path


def _replace_ellipsis(para, replacements: list[str]):
    """Заменяет … в прогоне параграфа значениями из списка по порядку."""
    i = 0
    for run in para.runs:
        while ELLIPSIS in run.text and i < len(replacements):
            run.text = run.text.replace(ELLIPSIS, replacements[i], 1)
            i += 1


def _fmt_cell(val: float) -> str:
    """Значение для 10-символьной колонки таблицы ПФ (правое выравнивание)."""
    if val == 0.0:
        return ' ' * _COL_W
    return f'{val:.4g}'.rjust(_COL_W)


def _make_table_row(num: float, den: float, deg: int, first: bool = False) -> str:
    """Формирует строку ASCII-таблицы CLASSiC с заданными значениями."""
    col1 = '| Ном.Система  |' if first else '|              |'
    return f"{col1}{_fmt_cell(num)} |{_fmt_cell(den)} |   {deg}   |"


def _rewrite_table_rows(rows: list, classic: dict):
    """Перезаписывает строки WP-таблицы (без «…»)."""
    num = classic.get('num', [0.0])
    den = classic.get('den', [0.0])

    def get(lst, i):
        return lst[i] if i < len(lst) else 0.0

    for row_i, para in enumerate(rows):
        if not para.runs:
            continue
        new_text = _make_table_row(get(num, row_i), get(den, row_i), row_i, row_i == 0)
        para.runs[0].text = new_text


def _fill_table_rows(paras: list, classic: dict):
    """
    Заполняет строки ASCII-таблицы ПФ коэффициентами.
    Вставляет дополнительные строки если степень > 2.
    """
    import copy

    num = classic.get('num', [0.0])
    den = classic.get('den', [0.0])

    def get(lst, i):
        return lst[i] if i < len(lst) else 0.0

    max_deg = max(len(num), len(den)) - 1

    def _write_row(run, deg, is_last=False):
        text = run.text
        n_val = _fmt_cell(get(num, deg))
        d_val = _fmt_cell(get(den, deg))
        text = text.replace(' ' * 9 + ELLIPSIS, n_val, 1)
        text = text.replace(' ' * 9 + ELLIPSIS, d_val, 1)
        if is_last:
            text = text.replace(f'   {ELLIPSIS}   ', f'   {deg}   ', 1)
        run.text = text

    if len(paras) > 0 and paras[0].runs:
        _write_row(paras[0].runs[0], 0)
    if len(paras) > 1 and paras[1].runs:
        _write_row(paras[1].runs[0], 1)

    last_para = paras[2] if len(paras) > 2 else None
    if last_para:
        for deg in range(2, max_deg):
            new_elem = copy.deepcopy(last_para._element)
            last_para._element.addprevious(new_elem)
            ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            for wt in new_elem.findall(f'.//{{{ns}}}t'):
                text = wt.text or ''
                n_val = _fmt_cell(get(num, deg))
                d_val = _fmt_cell(get(den, deg))
                text = text.replace(' ' * 9 + ELLIPSIS, n_val, 1)
                text = text.replace(' ' * 9 + ELLIPSIS, d_val, 1)
                text = text.replace(f'   {ELLIPSIS}   ', f'   {deg}   ', 1)
                wt.text = text
        if last_para.runs:
            _write_row(last_para.runs[0], max_deg, is_last=True)


def _insert_image_in_para(para, img_path: str, width_in: float = 5.5):
    """Добавляет изображение в существующий (пустой) параграф."""
    if not img_path or not Path(img_path).exists():
        return
    run = para.add_run()
    run.add_picture(img_path, width=Inches(width_in))


def _insert_image_before(doc, target_para, img_path: str, width_in: float = 5.5):
    """Вставляет новый параграф с изображением перед target_para."""
    if not img_path or not Path(img_path).exists():
        return
    tmp = doc.add_paragraph()
    tmp.alignment = 1  # CENTER
    tmp.add_run().add_picture(img_path, width=Inches(width_in))
    target_para._element.addprevious(tmp._element)


# ── Helpers для текстовой таблицы блоков ─────────────────────────────────────

def _fmt_blk_val(val: float) -> str:
    """10-символьное число с правым выравниванием для таблицы блоков.
    Ноль всегда отображается как '0' (CLASSiC показывает явный 0 в знаменателе интегратора).
    """
    return f'{val:.4g}'.rjust(_BLK_VAL_W)


def _blk_data_row(block_n: int, num: float, den: float, deg: int, conn: str) -> str:
    """Строка данных блока: |    #N      |  num  |  den  |  deg  | conn |"""
    c1 = f'    #{block_n}      '[:_BLK_ID_W]  # 12 символов, safe для #1-#9
    cc = f' {conn:<9}'                          # 10 символов: пробел + значение выравненное влево
    return f'|{c1}|{_fmt_blk_val(num)} |{_fmt_blk_val(den)} |   {deg}   |{cc}|'


def _blk_label_row(label: str) -> str:
    """Строка метки блока: | Вход       |           |           |       |          |"""
    c1 = f' {label:<{_BLK_ID_W - 1}}'  # 12 символов
    return (f'|{c1}|{" " * _BLK_VAL_W} |{" " * _BLK_VAL_W} '
            f'|       |          |')


def _blk_cont_row(den: float, deg: int) -> str:
    """Строка продолжения (deg > 0): |            |           |  den  |  deg  |          |"""
    return (f'|{" " * _BLK_ID_W}|{" " * _BLK_VAL_W} '
            f'|{_fmt_blk_val(den)} |   {deg}   |          |')


def _fill_block_table(paras: list, calc: CalcResults):
    """
    Перезаписывает строки текстовой таблицы блоков CLASSiC значениями варианта.

    Структура: block_params[N] = (num0, den0, deg0, conn, cont_den|None)
      cont_den — коэффициент den для строки deg=1 (T3 или T4), None если нет.
    """
    K1, K3, T3 = calc.K1, calc.K3, calc.T3
    K4, T4, K5 = calc.K4, calc.T4, calc.K5

    block_params = {
        1: (K1,  1.0, 0, '2',  None),
        2: (K3,  1.0, 0, '3',  T3),
        3: (K4,  1.0, 0, '4',  T4),
        4: (K5,  0.0, 0, '-1', None),
    }

    def _set(p, text):
        if p.runs:
            p.runs[0].text = text
            for r in p.runs[1:]:
                r.text = ''

    for i, p in enumerate(paras):
        m = _BLK_ROW_RE.match(p.text)
        if not m:
            continue
        bn = int(m.group(1))
        if bn not in block_params:
            continue
        num, den, deg, conn, cont_den = block_params[bn]
        _set(p, _blk_data_row(bn, num, den, deg, conn))
        if cont_den is not None and i + 1 < len(paras):
            nxt = paras[i + 1]
            if nxt.text.startswith('|            |'):
                _set(nxt, _blk_cont_row(cont_den, 1))


# ── Вычисление частотных показателей качества ────────────────────────────────

def _compute_freq_margins(calc: CalcResults) -> dict:
    """
    Вычисляет частоту среза, запас по фазе, частоту пи и запас по модулю
    из WP(s) численно (без scipy.signal.margin, только numpy).

    WP_classic хранит коэффициенты от s^0, scipy.poly1d ждёт от старшей степени.
    """
    try:
        import numpy as np

        wpc = calc.WP_classic
        if not wpc or not wpc.get('num') or not wpc.get('den'):
            return {}

        # Переворачиваем: WP_classic[s^0, s^1, ...] → poly1d[s^N, ..., s^0]
        num_poly = np.poly1d(list(reversed(wpc['num'])))
        den_poly = np.poly1d(list(reversed(wpc['den'])))

        # Частотная сетка: 10 000 точек от 1e-3 до 1e4 рад/с
        w = np.logspace(-3, 4, 10_000)
        H = num_poly(1j * w) / den_poly(1j * w)
        mag_db = 20.0 * np.log10(np.abs(H))
        # Разворачиваем фазу чтобы избежать разрывов при ±180°
        phase = np.degrees(np.unwrap(np.angle(H)))

        # ── Частота среза (wgc): mag пересекает 0 дБ снизу вверх (убывая) ──
        crossings = np.where(np.diff(np.sign(mag_db)))[0]
        wgc = pm = 0.0
        if len(crossings):
            idx = crossings[0]
            wgc = float(np.interp(0.0, [mag_db[idx + 1], mag_db[idx]],
                                       [w[idx + 1],       w[idx]]))
            phase_at_gc = float(np.interp(wgc, w, phase))
            pm = 180.0 + phase_at_gc

        # ── Частота пи (wpc): phase пересекает -180° ──
        shifted = phase + 180.0
        pc_crossings = np.where(np.diff(np.sign(shifted)))[0]
        wpc_freq = gm_db = 0.0
        if len(pc_crossings):
            idx = pc_crossings[0]
            wpc_freq = float(np.interp(0.0, [shifted[idx + 1], shifted[idx]],
                                            [w[idx + 1],         w[idx]]))
            mag_at_pc = float(np.interp(wpc_freq, w, mag_db))
            gm_db = -mag_at_pc  # запас по модулю: -L(ωπ) дБ

        return {
            'wgc':   wgc,
            'pm':    pm,
            'wpc':   wpc_freq,
            'gm_db': gm_db,
        }
    except Exception as e:
        print(f'[report] Предупреждение: частотные показатели не вычислены: {e}')
        return {}


# ── Выделение тестовых ответов ────────────────────────────────────────────────

def _bold_test_answers(paras: list):
    """
    Выделяет жирным правильные ответы на тестовые вопросы.

    Q3  (принцип управления): ответ «принцип замкнутого управления» —
        один длинный параграф со всеми вариантами; болдим нужный run.

    Q12 (область устойчивости при T2=0): параграф вида
        '1: (0K1.25);  2: (0K100);  3: (0K);  4: (K).'
        Ответ всегда вариант 3 — (0 < K < ∞); болдим runs от '3:' до '4:'.

    Q13 (АФЧХ из рис.7 для исходной системы): строка '1;\t2;\t3;\t4.'
        после вопроса «соответствует системе, заданной в задаче 2».
        Ответ всегда 3 (Тип 1, 2 апериодических звена).

    Q14 (АФЧХ из рис.7 при W5=K4): строка '1;\t2;\t3;\t4.'
        после вопроса «рис.7, соответствует такой системе».
        Ответ всегда 1 (Тип 0, 2 апериодических звена).

    Q15 (Найквист): параграф '3:система находится на колебательной границе...'
        Рис.8 одинаков для всех вариантов → ответ всегда 3.
    """
    pending_ris7 = None  # 13 or 14 — ожидаем строку '1;\t2;\t3;\t4.' для этого вопроса

    for p in paras:
        t = p.text.strip()

        # ── Трекинг контекста Q13/Q14 ──
        if 'из этих характеристик соответствует системе' in t and 'задаче 2' in t:
            pending_ris7 = 13
        elif 'рис.7, соответствует такой системе' in t:
            pending_ris7 = 14
        elif pending_ris7 and t == '1;\t2;\t3;\t4.':
            if pending_ris7 == 13 and len(p.runs) > 2:
                p.runs[2].bold = True   # '\t3;' — кривая 3
            elif pending_ris7 == 14 and len(p.runs) > 0:
                p.runs[0].bold = True   # '1;'  — кривая 1
            pending_ris7 = None

        # ── Q3: принцип замкнутого управления (обратная связь) ──
        if ('принцип разомкнутого управления' in t
                and 'принцип замкнутого управления' in t):
            for run in p.runs:
                if 'принцип замкнутого управления' in run.text:
                    run.bold = True

        # ── Q12: (0 < K < ∞) — вариант 3 всегда верен (система 2-го порядка, T3>0) ──
        # Параграф начинается с '1: (0' и заканчивается 'K).'
        elif t.startswith('1: (0') and t.endswith('K).'):
            in_opt3 = False
            for run in p.runs:
                if '4:' in run.text:
                    break
                if '3:' in run.text:
                    in_opt3 = True
                if in_opt3:
                    run.bold = True

        # ── Q15: колебательная граница (без пробела после '3:' — отличает от Q10) ──
        elif t == '3:система находится на колебательной границе устойчивости,':
            for run in p.runs:
                run.bold = True


# ── Main ──────────────────────────────────────────────────────────────────────

def fill_report(
    template_path: str,
    output_dir:    str,
    calc:          CalcResults,
    shots:         Screenshots,
) -> str:
    """
    Заполняет задание LP_vXX результатами расчётов и скриншотами.
    Возвращает путь к готовому .docx файлу.
    """
    src = Path(template_path)
    out_dir = Path(output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    variant_str = f'{calc.variant:02d}'
    out_path = out_dir / f'LP_v{variant_str}_done.docx'

    # 1. .doc → .docx если нужно
    if src.suffix.lower() == '.doc':
        docx_src = src.with_suffix('.docx')
        if not docx_src.exists():
            print('[report] Конвертация .doc → .docx...')
            _convert_doc_to_docx(str(src), str(docx_src))
        src = docx_src

    shutil.copy(src, out_path)
    doc = Document(str(out_path))
    paras = doc.paragraphs

    # Извлекаем только правую часть формул (убираем "X(s) = " префикс)
    def _rhs(formula: str) -> str:
        return formula.split(' = ', 1)[-1] if ' = ' in formula else formula

    wp_rhs   = _rhs(calc.WP_str)
    phi_rhs  = _rhs(calc.Phi_str)
    phie_rhs = _rhs(calc.PhiE_str)

    # Предварительно вычисляем частотные показатели
    margins = _compute_freq_margins(calc)

    # 2. ── WP-формула и WP-таблица ────────────────────────────────────────
    for p in paras:
        txt = p.text
        if txt.startswith('WP(s)=') and ELLIPSIS not in txt and len(txt) > 7:
            for i, run in enumerate(p.runs):
                if '= ' in run.text or run.text == '=':
                    run.text = f'= {wp_rhs}'
                    for r2 in p.runs[i + 1:]:
                        r2.text = ''
                    break
            break

    wp_rows: list = []
    for p in paras:
        if p.text.startswith('|') and ELLIPSIS in p.text:
            break
        if ELLIPSIS not in p.text and re.search(r'\|   \d+   \|$', p.text):
            wp_rows.append(p)

    _rewrite_table_rows(wp_rows, calc.WP_classic)

    # 3. ── Текстовые замены и поиск подписей к рисункам ──────────────────
    phie_count = 0
    ris1a = ris3 = ris4 = ris5 = ris6 = None

    for p in paras:
        txt = p.text
        t   = txt.strip()

        # ── Поиск подписей к рисункам ──
        if 'Рис.1а' in t or 'Рис. 1а' in t:
            ris1a = p
        if t == 'Рис. 3':
            ris3 = p
        elif t == 'Рис. 4':
            ris4 = p
        elif 'Рис. 5' in t:
            ris5 = p
        elif t == 'Рис. 6':
            ris6 = p

        # ── Замены ──
        if txt.startswith('Ф(s)=') and ELLIPSIS in txt:
            _replace_ellipsis(p, [phi_rhs])

        elif txt.startswith('Фe(s)=') and ELLIPSIS in txt:
            phie_count += 1
            if phie_count == 1:
                _replace_ellipsis(p, ['1 / (1 + WP(s))'])
            elif phie_count == 2:
                _replace_ellipsis(p, [phie_rhs])

        elif txt.startswith('eуст=lim') and ELLIPSIS in txt:
            _replace_ellipsis(p, ['Фe(0)', calc.e_ust_step])

        elif txt.startswith('eуст=') and ELLIPSIS in txt and 'lim' not in txt:
            _replace_ellipsis(p, [calc.e_ust_ramp])

        elif txt.startswith('Kкр=') and ELLIPSIS in txt:
            _replace_ellipsis(p, [str(round(calc.K_loop_critical, 4))])

        elif 'Модель сохранена в файле' in txt and ELLIPSIS in txt:
            for j, run in enumerate(p.runs):
                if ELLIPSIS in run.text:
                    run.text = run.text.replace(ELLIPSIS, f'var{variant_str}')
                    if j + 1 < len(p.runs):
                        p.runs[j + 1].text = p.runs[j + 1].text.lstrip()
                    break

        # Имя модели в текстовой форме: 'Модель: "….MDL"'
        elif txt.startswith('Модель: "') and ELLIPSIS in txt:
            for run in p.runs:
                if ELLIPSIS in run.text:
                    run.text = run.text.replace(ELLIPSIS, f'VAR{variant_str}')
                    break

        # Частотные показатели качества (задача 13)
        elif margins and txt.startswith('\tЧастота среза:') and p.runs:
            p.runs[0].text = f'\tЧастота среза: {margins["wgc"]:.4f} рад/с'
            for r in p.runs[1:]:
                r.text = ''

        elif margins and txt.startswith('\tЗапас по фазе:') and p.runs:
            p.runs[0].text = f'\tЗапас по фазе: {margins["pm"]:.4f} град'
            for r in p.runs[1:]:
                r.text = ''

        elif margins and txt.startswith('\tЧастота пи:') and p.runs:
            p.runs[0].text = f'\tЧастота пи: {margins["wpc"]:.4f} рад/с'
            for r in p.runs[1:]:
                r.text = ''

        elif margins and txt.startswith('\tЗапас по модулю:') and p.runs:
            p.runs[0].text = f'\tЗапас по модулю: {margins["gm_db"]:.4f} дБ'
            for r in p.runs[1:]:
                r.text = ''

    # 4. ── Таблицы коэффициентов ПФ ───────────────────────────────────────
    table_groups: list[list] = []
    buf: list = []
    for p in paras:
        if p.text.startswith('|') and ELLIPSIS in p.text:
            buf.append(p)
            if len(buf) == 3:
                table_groups.append(buf)
                buf = []
        else:
            buf = []

    if len(table_groups) >= 1:
        _fill_table_rows(table_groups[0], calc.Phi_classic)
    if len(table_groups) >= 2:
        _fill_table_rows(table_groups[1], calc.PhiE_classic)

    # 5. ── Текстовая таблица блоков ───────────────────────────────────────
    if calc.K1 != 0.0:
        _fill_block_table(paras, calc)

    # 6. ── Тестовые вопросы — выделение правильных ответов жирным ───────────

    # Q10: устойчивость по Гурвицу
    if not calc.hurwitz_stable:
        coeffs = calc.char_coeffs
        if all(c > 0 for c in coeffs):
            q10_answer = '2:  система нейтральна (находится на нейтральной границе устойчивости),'
        else:
            q10_answer = '4:  система неустойчива.'
    else:
        q10_answer = '1:  система устойчива,'

    for p in paras:
        if p.text.strip() == q10_answer.strip():
            for run in p.runs:
                run.bold = True
            break

    # Q3, Q12, Q15: принцип управления, область устойчивости, Найквист
    _bold_test_answers(paras)

    # 7. ── Скриншоты ──────────────────────────────────────────────────────
    # Рис.1а: структурная схема (вставляем перед подписью)
    if ris1a and shots.schema:
        _insert_image_before(doc, ris1a, shots.schema)

    # Рис.3: пустой параграф перед подписью
    if ris3:
        idx = paras.index(ris3)
        if idx > 0 and paras[idx - 1].text == '':
            _insert_image_in_para(paras[idx - 1], shots.step_response)

    # Рис.4: рамповый вход f(t)=0.1t — вставляем перед подписью
    if ris4 and shots.ramp_response:
        _insert_image_before(doc, ris4, shots.ramp_response)

    # Рис.5: вставляем перед подписью
    if ris5 and shots.critical:
        _insert_image_before(doc, ris5, shots.critical)

    # Рис.6: вставляем перед подписью
    if ris6 and shots.bode:
        _insert_image_before(doc, ris6, shots.bode)

    doc.save(str(out_path))
    print(f'[report] Готово: {out_path}')
    return str(out_path)
